# Chapter 15: Woking with PDF and word documents

import PyPDF2
import docx 


# 1. Opening and Reading PDF Files
def open_pdf(file_path):
    try:
        pdf_file = open(file_path, 'rb')
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        print(f"PDF '{file_path}' opened successfully.")
        return pdf_reader
    except FileNotFoundError:
        print(f"Error: PDF file '{file_path}' not found.")
        return None

# 2. Extracting Text from a PDF
def extract_pdf_text(pdf_reader, page_number):
    try:
        page = pdf_reader.pages[page_number]
        text = page.extract_text()
        print(f"Text from page {page_number}:\n{text}")
        return text
    except IndexError:
        print(f"Error: Page {page_number} does not exist in the PDF.")
        return None

# 3. Extracting All Text from a PDF
def extract_all_text(pdf_reader):
    all_text = ""
    for page_num in range(len(pdf_reader.pages)):
        all_text += pdf_reader.pages[page_num].extract_text() + "\n"
    print("Extracted all text from PDF.", all_text)
    return all_text

# 4. Combining Pages from Multiple PDFs
# def combine_pdfs(pdf_list, output_path):
#     pdf_writer = PyPDF2.PdfWriter()
#     try:
#         for pdf_path in pdf_list:
#             pdf_reader = PyPDF2.PdfReader(open(pdf_path, 'rb'))
#             for page_num in range(len(pdf_reader.pages)):
#                 pdf_writer.add_page(pdf_reader.pages[page_num])
#         with open(output_path, 'wb') as output_pdf:
#             pdf_writer.write(output_pdf)
#         print(f"Combined PDFs saved to '{output_path}'.")
#     except Exception as e:
#         print(f"Error combining PDFs: {e}")


# 5. Creating a New Word Document
def create_word_doc(file_path):
    doc = docx.Document()
    doc.add_heading('Document Title', 0)
    doc.save(file_path)
    print(f"New Word document '{file_path}' created.")

# 6. Adding Text to a Word Document
def add_text_to_word(file_path, text, style='Normal'):
    try:
        doc = docx.Document(file_path)
        doc.add_paragraph(text, style=style)
        doc.save(file_path)
        print(f"Text added to '{file_path}' with style '{style}'.")
    except Exception as e:
        print(f"Error adding text to Word document: {e}")

# 7. Reading Text from a Word Document
def read_word_doc(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        combined_text = '\n'.join(full_text)
        print(f"Content of '{file_path}':\n{combined_text}")
        return combined_text
    except Exception as e:
        print(f"Error reading Word document: {e}")
        return None

# 8. Adding Tables to a Word Document
def add_table_to_word(file_path, data):
    try:
        doc = docx.Document(file_path)
        table = doc.add_table(rows=1, cols=len(data[0]))
        
        # Add header row
        hdr_cells = table.rows[0].cells
        for i, heading in enumerate(data[0]):
            hdr_cells[i].text = heading
        
        # Add the rest of the data
        for row_data in data[1:]:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        doc.save(file_path)
        print(f"Table added to '{file_path}'.")
    except Exception as e:
        print(f"Error adding table to Word document: {e}")
